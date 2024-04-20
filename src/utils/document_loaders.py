import docx
import fitz
from typing import List

from utils.documents import Document


class DocumentLoader:
    """Load documents from various file formats: .txt, .pdf, .doc, .docx."""

    def __init__(self, document_path: str) -> None:
        """
        Create a new DocumentLoader.

        Args:
            document_path (str): Path to document.
        """

        self.document_path = document_path

    def load(self) -> List[Document]:
        """
        Load documents from file.

        Returns:
            List[Document]: List of documents.
        """

        documents = []

        if self.document_path.endswith(".pdf"):
            doc = fitz.open(self.document_path)
            for i, page in enumerate(doc):
                text = page.get_text()
                metadata = {"source": self.document_path, "page": i}
                documents.append(Document(page_content=text, metadata=metadata))
        elif self.document_path.endswith(".txt"):
            with open(self.document_path, "r", encoding="utf-8") as txt_file:
                text = txt_file.read()
                metadata = {"source": self.document_path}
                documents.append(Document(page_content=text, metadata=metadata))
        elif self.document_path.endswith((".doc", ".docx")):
            doc = docx.Document(self.document_path)
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text
                metadata = {"source": self.document_path, "page": i}
                documents.append(Document(page_content=text, metadata=metadata))
        else:
            print(f"Unsupported file format: {self.document_path}")

        return documents
