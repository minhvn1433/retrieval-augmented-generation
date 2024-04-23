import gc
import time

import docx
import fitz
import torch
from sentence_transformers import util

from globals import bi_encoder, cross_encoder, tokenizer, llm, chat
from utils.document import Document
from utils.text_splitters import CharacterTextSplitter, RecursiveCharacterTextSplitter


def clear_cuda_cache():
    torch.cuda.empty_cache()
    gc.collect()


def load_documents(files):
    docs = []

    for file in files:
        if file.endswith(".pdf"):
            doc = fitz.open(file)
            for i, page in enumerate(doc):
                text = page.get_text()
                metadata = {"source": file, "page": i}
                docs.append(Document(page_content=text, metadata=metadata))
        elif file.endswith(".txt"):
            with open(file, "r", encoding="utf-8") as txt_file:
                text = txt_file.read()
                metadata = {"source": file}
                docs.append(Document(page_content=text, metadata=metadata))
        elif file.endswith((".doc", ".docx")):
            doc = docx.Document(file)
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text
                metadata = {"source": file, "page": i}
                docs.append(Document(page_content=text, metadata=metadata))
        else:
            print(f"Unsupported file format: {file}")

    return docs


def reordering_texts(texts):
    texts.reverse()
    reordered_result = []
    for i, value in enumerate(texts):
        if i % 2 == 1:
            reordered_result.append(value)
        else:
            reordered_result.insert(0, value)
    return reordered_result


def user(message, history):
    return "", history + [[message, None]]


def index_files(files):
    global split_texts, corpus_embeddings

    # 1. LOAD
    docs = load_documents(files)

    # 2. SPLIT
    text_splitter = CharacterTextSplitter(chunk_size=512, chunk_overlap=256)
    split_docs = text_splitter.split_documents(docs)

    # 3. EMBED
    split_texts = [doc.page_content for doc in split_docs]
    corpus_embeddings = bi_encoder.encode(split_texts, convert_to_tensor=True)

    # 4. STORE

    return "Finished!"


def bot(history):
    # 5. RETRIEVE
    query = history[-1][0]
    query_embedding = bi_encoder.encode(query, convert_to_tensor=True)
    hits = util.semantic_search(query_embedding, corpus_embeddings, top_k=20)[0]
    indices = [hit["corpus_id"] for hit in hits]
    unique_texts = [split_texts[i] for i in indices]

    # 6. RE-RANK
    cross_scores = cross_encoder.predict([(query, text) for text in unique_texts])
    reranked_texts = [
        text for _, text in sorted(zip(cross_scores, unique_texts), reverse=True)[:5]
    ]
    reordered_texts = reordering_texts(reranked_texts)

    # 7. AUGMENT
    context = "\n\n".join(reordered_texts)
    augmented_prompt = f"""You are an assistant for question-answering tasks. \
Use the following pieces of retrieved context to answer the question. If you \
don't know the answer, just say that you don't know.

Context: 
{context}
        
Question: 
{query}"""

    # 8. GENERATE
    chat.append({"role": "user", "content": augmented_prompt})
    input_text = tokenizer.apply_chat_template(
        chat, tokenize=False, add_generation_prompt=True
    )
    inputs = tokenizer.encode(
        input_text, add_special_tokens=False, return_tensors="pt"
    ).to(llm.device)
    outputs = llm.generate(
        inputs,
        temperature=0.1,
        do_sample=True,
        max_new_tokens=1024,
        repetition_penalty=1.1,
    )
    output_text = tokenizer.decode(outputs[0])
    response = (
        output_text.replace(input_text, "").replace("<bos>", "").replace("<eos>", "")
    )
    chat.append({"role": "assistant", "content": response})

    # 9. STREAM
    history[-1][1] = ""
    for character in response:
        history[-1][1] += character
        time.sleep(0.01)
        yield history
